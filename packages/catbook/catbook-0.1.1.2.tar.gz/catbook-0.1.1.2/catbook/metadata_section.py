from . import Markup
from . import Fonts
from . import BookMetadata
from . import Metadata
from . import Section
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import traceback
from typing import List, Optional
from . import Tokenizer
from datetime import datetime
from typing import cast


class MetadataSection(Section):
    def __init__(
        self,
        lines: List[str],
        markup: Markup,
        fonts: Fonts,
        document: Document,
        metadata: Metadata,
    ) -> None:
        super().__init__(lines, markup, fonts, document, metadata)

    @property
    def metadata(self) -> BookMetadata:
        return cast(BookMetadata, self._metadata)

    # ============= PUBLIC STUFF HERE

    def compile(self) -> bool:
        """returns True if the lines were all added to the document"""
        p = self.doc.add_paragraph()
        run = self._add_run(p, "")
        run.font.name = self._fonts.BODY
        run.add_break(WD_BREAK.PAGE)

        run = self._add_run(p, "Book metadata", 12, True)

        table = self.doc.add_table(rows=1, cols=2)
        table.style.name = "Table Grid"

        row = table.rows[0].cells
        p = row[0].paragraphs[0]
        self._add_run(p, "Title", 10, False)
        p = row[1].paragraphs[0]
        self._add_run(p, f"{self.metadata.TITLE}", 10, False)

        row = table.add_row().cells
        p = row[0].paragraphs[0]
        self._add_run(p, "Author", 10, False)
        p = row[1].paragraphs[0]
        self._add_run(p, f"{self.metadata.AUTHOR}", 10, False)

        row = table.add_row().cells
        p = row[0].paragraphs[0]
        self._add_run(p, "Compiled on", 10, False)
        p = row[1].paragraphs[0]
        self._add_run(p, f"{datetime.now()}", 10, False)

        row = table.add_row().cells
        p = row[0].paragraphs[0]
        self._add_run(p, "Compiled from", 10, False)
        p = row[1].paragraphs[0]
        self._add_run(p, f"{self.metadata.FILES.INPUT}", 10, False)

        row = table.add_row().cells
        p = row[0].paragraphs[0]
        self._add_run(p, "Total sections", 10, False)
        p = row[1].paragraphs[0]
        self._add_run(p, f"{self.metadata.count}", 10, False)

        row = table.add_row().cells
        p = row[0].paragraphs[0]
        self._add_run(p, "Stand alone sections", 10, False)
        p = row[1].paragraphs[0]
        self._add_run(p, f"{self.metadata.stand_alone_count}", 10, False)

        row = table.add_row().cells
        p = row[0].paragraphs[0]
        self._add_run(p, "Words", 10, False)
        p = row[1].paragraphs[0]
        self._add_run(p, f"{self.metadata.word_count}", 10, False)

        row = table.add_row().cells
        p = row[0].paragraphs[0]
        self._add_run(p, "Unique words", 10, False)
        p = row[1].paragraphs[0]
        self._add_run(p, f"{self.metadata.unique_words_count()}", 10, False)

        p = self.doc.add_paragraph()
        run = self._add_run(p, "")
        p = self.doc.add_paragraph()
        run = self._add_run(p, "Compiled by catbook", 8, False, True)

        return True

    # ============= INTERNAL STUFF STARTS HERE

    def _add_run(
        self,
        p: Paragraph,
        text: str,
        size: int = 10,
        bold: bool = False,
        italic: bool = False,
    ) -> Run:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = self._fonts.BLOCK
        run.font.size = Pt(size)
        return run
