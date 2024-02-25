from . import Markup
from . import Fonts
from . import SectionMetadata
from . import Metadata
from . import Section
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import traceback
from typing import List, Optional, cast
from . import Tokenizer


class RegularSection(Section):
    def __init__(
        self,
        lines: List[str],
        markup: Markup,
        fonts: Fonts,
        document: Document,
        metadata: Metadata,
    ) -> None:
        super().__init__(lines, markup, fonts, document, metadata)

    # ============= PUBLIC STUFF HERE

    def compile(self) -> bool:
        """returns True if the lines were all added to the document"""
        reqs = [self._markup, self._fonts, self._document, self._metadata]
        if None in reqs:
            raise Exception(f"Requirements not setup: {reqs}")
        line_number = 0
        for line in self._lines:
            try:
                line_number = self._append_line(self._lines, line, line_number)
                line_number = line_number + 1
                self.metadata.FILE_LINE_COUNT = line_number
            except Exception as e:
                print(f"Error in section at content line {line_number}: {line}: {e}")
                return False
        return True

    @property
    def doc(self) -> Document:
        return self._document

    @property
    def metadata(self) -> SectionMetadata:
        return cast(SectionMetadata, self._metadata)

    # ============= INTERNAL STUFF STARTS HERE

    def _add_paragraph(self) -> Paragraph:
        p = self.doc.add_paragraph()
        # in testing we might not have a self.metadata, so being defensive.
        if self.metadata:
            self.metadata.PARAGRAPH_COUNT = self.metadata.PARAGRAPH_COUNT + 1
        return p

    def _append_quote(self) -> None:
        """adds the quote into the document"""
        if self._quote is None:  # type: ignore [has-type]
            raise Exception("No quote array found at _append_quote")
        i = len(self._quote)
        for aline in self._quote:
            if aline is None:  # make mypy happy
                continue
            p = self._add_paragraph()
            paragraph_format = p.paragraph_format
            run = self._add_run(p, f"   {aline[1:]}")
            run.font.name = self._fonts.QUOTE
            run.italic = True
            i = i - 1
            if i > 0:
                paragraph_format.space_after = Pt(1)
        self._quote = None

    def _append_block(self) -> None:
        """adds the block text to the document"""
        if not self._block:
            raise Exception("No block array found at _append_block")
        i = len(self._block)
        for aline in self._block:
            if aline is None:  # make mypy happy
                continue
            p = self._add_paragraph()
            paragraph_format = p.paragraph_format
            thisline = aline[1:]
            run = self._add_run(p, f"{thisline}")
            run.italic = True
            run.font.name = self._fonts.BLOCK
            run.font.size = Pt(10)
            i = i - 1
            if i > 0:
                paragraph_format.space_after = Pt(1)
        self._block = None

    def _append_output(self, lines: List[str], line: str, line_number: int) -> None:
        """writes blocks, quotes, and page breaks to the document"""
        if self._block is not None:
            # write the block
            self._append_block()
        if self._quote is not None:
            # write the quote
            self._append_quote()
        if self._part_break and self._last_line(lines, line_number):
            #
            # we never get here?!
            #
            """
            p = self._add_paragraph()
            run = self._add_run(p, "")
            run.font.name = self._fonts.BODY
            run.add_break(WD_BREAK.PAGE)
            self._part_break = False
            self._book_break = False
            self._last_was_break = True
            print("setting last_was_break=True")
            """

    def _last_line(self, lines: List[str], line_number: int) -> bool:
        """returns True if line_number indicates the last non-blank line"""
        n = len(lines)
        if line_number == n:
            return True
        for r in range(line_number + 1, n):
            if lines[r].strip() != "":
                return False
        return True

    # ====================
    # start title stuff
    # ====================

    def _append_title(self, line: str) -> None:
        if self._add_jump_if(line):
            return
        self._consider_title(line)
        self._handle_chapter_if(line)
        self._handle_book_if(line)
        self._handle_section_if(line)

    def _consider_title(self, line: str) -> None:
        self._is_chapter(line)
        self._is_book(line)
        self._is_section(line)

    def _add_page_break_if(self) -> None:
        c = self.metadata.CHAPTER
        b = self.metadata.BOOK
        lwb = self._last_was_break
        s = self.metadata.NEW_SECTION
        cnt = self.metadata.BOOK_METADATA.paragraphs_count()  # type: ignore
        if self._needs_page_break_before(
            chapter=c, book=b, last_was_break=lwb, new_section=s, paragraph_count=cnt
        ):
            p = self._add_paragraph()
            run = self._add_run(p, "")
            run.font.name = self._fonts.TITLE
            run.add_break(WD_BREAK.PAGE)
            self.metadata.STAND_ALONE = True
            return p

    def _needs_page_break_before(
        self,
        *,
        chapter: bool,
        book: bool,
        last_was_break: bool,
        new_section: bool,
        paragraph_count: int,
    ) -> bool:
        return (
            (chapter and not last_was_break)
            or (book and not last_was_break)
            or new_section
        ) and paragraph_count > 0

    def _handle_chapter_if(self, line) -> None:
        if self._is_chapter(line):
            self.metadata.CHAPTER = True
            line = line[1:]
            self._add_upper_heading(line, 1)

    def _handle_book_if(self, line) -> None:
        if self._is_book(line):
            self.metadata.BOOK = True
            line = line[2:]
            self._add_upper_heading(line, 1)

    def _add_upper_heading(self, line: str, level: int) -> None:
        self._add_page_break_if()
        p = self.doc.add_heading("", level)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = self._add_run(p, line)
        run.font.name = self._fonts.BODY
        run.font.color.rgb = RGBColor.from_string("000000")
        self.metadata.PARAGRAPH_COUNT = self.metadata.PARAGRAPH_COUNT + 1

    def _handle_section_if(self, line) -> None:
        if self._is_section(line):
            if self._is_new_section(line):
                self.metadata.NEW_SECTION = True
                while line.find(self.markup.NEW_SECTION) == 0:
                    line = line[1:]
                self._add_page_break_if()
            p = self.doc.add_heading("", 3)
            paragraph_format = p.paragraph_format
            paragraph_format.space_before = Pt(30)
            paragraph_format.space_after = Pt(10)
            run = self._add_run(p, line)
            run.font.name = self._fonts.BODY
            run.font.color.rgb = RGBColor.from_string("000000")
            self.metadata.PARAGRAPH_COUNT = self.metadata.PARAGRAPH_COUNT + 1

    def _is_new_section(self, line: str) -> bool:
        return (len(line) >= 1 and line[0:1] == self._markup.NEW_SECTION) or (
            self.metadata.last_section() is not None
            and cast(SectionMetadata, self.metadata.last_section()).is_book_or_chapter()
        )  # type: ignore

    def _is_chapter(self, line: str) -> bool:
        if len(line) == 0:
            return False
        return line[0] == self._markup.CHAPTER_TITLE and not self._is_book(line)

    def _is_book(self, line: str) -> bool:
        return len(line) >= 2 and line[0:2] == self._markup.BOOK_TITLE

    def _is_section(self, line: str) -> bool:
        return not self._is_book(line) and not self._is_chapter(line)

    # ====================
    # end of title stuff
    # ====================

    def _handle_block(self, line, line_number, lines: int) -> Optional[bool]:
        block = self._markup._is_block(line, line_number, lines)
        if block:
            if self._block is None:
                self._block = []
            self._block.append(line)
            # pack the line into self.block
            return True
        elif block is None:
            return None
        elif self._block is not None:
            # write the block
            self._append_block()
            return False
        else:
            return False

    def _handle_quote(self, line: str, line_number: int, lines: int) -> Optional[bool]:
        quote = self._markup._is_quote(line, line_number, lines)
        if quote:
            if self._quote is None:
                self._quote = []
            self._quote.append(line)
            # we pack the line into self.quote
            return True
        elif self._quote is not None:
            # write the quote
            self._append_quote()
            return False
        else:
            return False

    def _handle_highlights(self, line):
        """adds a line that includes a highlight to the document"""
        block = 0
        p = self._add_paragraph()
        while self._markup.WORD_HIGHLIGHT in line:
            start = line.index(self._markup.WORD_HIGHLIGHT)
            end = line.index(self._markup.WORD_HIGHLIGHT, start + 1)
            front = line[0:start]
            mid = line[start + 1 : end]
            back = line[end + 1 :]
            line = back
            run = self._add_run(p, "   " if block == 0 else "")
            run.font.name = self._fonts.BODY
            run = self._add_run(p, front)
            run.font.name = self._fonts.BODY
            run = self._add_run(p, mid)
            run.italic = True
            run.font.name = self._fonts.BODY
            if self._markup.WORD_HIGHLIGHT not in line:
                run = self._add_run(p, back)
                run.font.name = self._fonts.BODY
            block = block + 1

    def _add_run(self, p: Paragraph, text: str) -> Run:
        if self.metadata is not None:
            ws = self._get_words(text)
            self.metadata.WORDS = ws + self.metadata.WORDS
            n = len(ws)
            self.metadata.WORD_COUNT = self.metadata.WORD_COUNT + n
        run = p.add_run(text)
        return run

    def _get_words(self, text: str) -> List[str]:
        words = Tokenizer.get_words(text)
        return words

    def _count_words(self, text: str) -> int:
        return len(self._get_words(text))

    def _add_jump_if(self, line: str) -> Paragraph:
        simple_separator = len(line) >= 3 and line[0:3] == self._markup.JUMP
        if simple_separator:
            self.metadata.JUMP = True
            p = self._add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph_format = p.paragraph_format
            paragraph_format.space_before = Pt(24)
            paragraph_format.space_after = Pt(25)
            run = p.add_run(self.markup.ASTERISM)
            run.bold = True
            return p
        return False

    def _handle_comment_if(self, lines: List[str], line: str) -> bool:
        if len(line) > 0 and line[0] == "#":
            self._include_image_if(lines, line)
            self._include_mark_if(lines, line)
            self._dump_metadata_if(lines, line)
            return True
        return False

    def _include_mark_if(self, lines: List[str], line: str) -> bool:
        token = "MARK"
        mark = line.find(token)
        if mark > 0:
            p = self._add_paragraph()
            run = p.add_run(
                f"file: {self.metadata.FILE}:{self.metadata.FILE_LINE_COUNT}"
            )
            run.font.name = self._fonts.QUOTE
            run.italic = True
            run.font.color.rgb = RGBColor.from_string("555555")
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        return False

    def _dump_metadata_if(self, lines: List[str], line: str) -> bool:
        token = "METADATA"
        mark = line.find(token)
        if mark > 0:
            p = self._add_paragraph()
            run = p.add_run(f"{self.metadata}")
            run.font.name = self._fonts.QUOTE
            run.italic = True
            run.font.color.rgb = RGBColor.from_string("555555")
            run.font.size = Pt(9)
            return True
        return False

    def _include_image_if(self, lines: List[str], line: str) -> bool:
        token = "INCLUDE IMAGE:"
        insert = line.find(token)
        if insert > 0:
            path = line[insert + len(token) :].strip()
            p = self._add_paragraph()
            run = p.add_run()
            run.add_picture(path)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        return False

    def _append_line(self, lines: List[str], line: str, line_number: int):
        try:
            line = line.strip()
            #
            # blank lines are paragraph breaks. we take this time
            # to write out blocks, quotes, etc.
            #
            if line == "":
                self._append_output(lines, line, line_number)
                return line_number
            #
            # comments are lines starting with #. they do not
            # count for the purposes of finding the title line.
            #
            if self._handle_comment_if(lines, line):
                return line_number - 1
            #
            # titles
            #
            if line_number == 0:
                self.metadata.set_first_line(self._markup, line)
                self._append_title(line)
                return line_number
            #
            # blocks
            #
            block = self._handle_block(line, line_number, len(lines))
            if block is None:
                # we found an escaped pipe: || meaning an
                # italicized word started a line
                line = line[1:]
            elif block:
                return line_number
            #
            # quotes
            #
            quote = self._handle_quote(line, line_number, len(lines))
            if quote:
                return line_number
            #
            # regular line
            #
            if self._markup.WORD_HIGHLIGHT in line:
                self._handle_highlights(line)
            else:
                p = self._add_paragraph()
                run = self._add_run(p, f"   {line}")
                run.font.name = self._fonts.BODY
            # self._last_was_break = False
            return line_number
        except Exception as e:
            print(f"Section failed at _append_line: line: {line_number}. error: {e}")
            traceback.print_exc()
